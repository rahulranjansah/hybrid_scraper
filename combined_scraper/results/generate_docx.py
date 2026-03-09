import csv
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

CSV_PATH = Path(__file__).parent / "hr_candidates_round2.csv"
DOCX_PATH = Path(__file__).parent / "HR_Candidates_Japan_Round2.docx"

# Japanese name -> English (romaji) translations
JP_NAME_MAP = {
    "八木洋介": "Yosuke Yagi",
    "有沢 正人": "Masato Arisawa",
    "西田 政之": "Masayuki Nishida",
    "木下 達夫": "Tatsuo Kinoshita",
    "木下達夫": "Tatsuo Kinoshita",
    "向井 麗子": "Reiko Mukai",
    "橋本佳介": "Keisuke Hashimoto",
    "石川 昇": "Noboru Ishikawa",
    "秋岡 和寿": "Kazutoshi Akioka",
    "鬼頭 伸彰": "Nobuaki Kito",
    "村田宗一郎": "Soichiro Murata",
    "村田 宗一郎": "Soichiro Murata",
    "有賀 誠": "Makoto Ariga",
    "清家良太": "Ryota Seike",
    "丸林哲也": "Tetsuya Marubayashi",
    "源田": "Genda",
    "佐藤": "Sato",
}

FLAG_COLORS = {
    "green": RGBColor(0x22, 0x8B, 0x22),
    "yellow": RGBColor(0xCC, 0x88, 0x00),
    "red": RGBColor(0xCC, 0x00, 0x00),
}


def clean_html(text: str) -> str:
    return re.sub(r"<[^>]+>", "", text).strip()


def translate_name(name: str) -> str:
    return JP_NAME_MAP.get(name, name)


def split_multi(value: str) -> list[str]:
    return [v.strip() for v in value.split(";") if v.strip()]


def extract_linkedin(url: str, linkedin_col: str) -> str:
    """Get best LinkedIn URL from source URL or linkedin column."""
    # Prefer linkedin column if it's a proper profile link
    for li in linkedin_col.split(";"):
        li = li.strip()
        if li and "/in/" in li:
            return li
    # Fall back to source URL if it's LinkedIn
    if "linkedin.com/in/" in url:
        return url
    return ""


def build_relevant_info(title: str, company: str, reason: str) -> str:
    parts = []
    if title:
        parts.append(title)
    if company:
        parts.append(f"at {company}")
    info = ", ".join(parts) if parts else ""
    if reason:
        info += f". {reason}" if info else reason
    return info


def load_candidates(path: Path) -> list[dict]:
    candidates = []
    seen_names = set()

    with open(path, newline="", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        for row in reader:
            score = int(row["relevance_score"])
            names = row.get("people_names", "").strip()
            if score == 0 or not names:
                continue

            name_list = split_multi(names)
            title_list = split_multi(row.get("people_titles", ""))
            company_list = split_multi(row.get("people_companies", ""))
            linkedin_col = row.get("people_linkedin", "")
            url = row.get("url", "")
            flag = row.get("flag", "")
            reason = clean_html(row.get("score_reason", ""))

            for i, name in enumerate(name_list):
                eng_name = translate_name(clean_html(name))

                # Deduplicate by English name
                key = eng_name.lower().strip()
                if key in seen_names:
                    continue
                seen_names.add(key)

                title = clean_html(title_list[i]) if i < len(title_list) else ""
                company = clean_html(company_list[i]) if i < len(company_list) else ""
                linkedin = extract_linkedin(url, linkedin_col)

                candidates.append({
                    "name": eng_name,
                    "linkedin": linkedin,
                    "info": build_relevant_info(title, company, reason),
                    "score": score,
                    "flag": flag,
                })

    candidates.sort(key=lambda c: c["score"], reverse=True)
    return candidates


def set_cell_text(cell, text, bold=False, size=Pt(9),
                  alignment=None, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if alignment:
        p.alignment = alignment
    run = p.add_run(str(text))
    run.font.size = size
    run.font.name = "Calibri"
    run.bold = bold
    if color:
        run.font.color.rgb = color


def add_hyperlink(paragraph, url, text, font_size=Pt(8)):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(font_size.pt * 2)))
    rPr.append(sz)

    font = OxmlElement("w:rFonts")
    font.set(qn("w:ascii"), "Calibri")
    rPr.append(font)

    run.append(rPr)
    run.text = text
    hyperlink.append(run)
    paragraph._element.append(hyperlink)


def build_docx(candidates: list[dict], output: Path):
    doc = Document()

    title = doc.add_heading(
        "HR Leadership Candidates — Japan (Round 2)", level=1
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "Sourced candidates for senior HR roles "
        "(CHRO, HR Director, HRBP, Country HR Manager) in Japan.\n"
        f"Total new candidates identified: {len(candidates)}"
    )

    headers = ["#", "Name", "LinkedIn", "Relevant Info", "Flag"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, header in enumerate(headers):
        set_cell_text(
            table.rows[0].cells[i], header, bold=True,
            size=Pt(9), alignment=WD_ALIGN_PARAGRAPH.CENTER,
        )

    for idx, c in enumerate(candidates, start=1):
        row = table.add_row()
        set_cell_text(
            row.cells[0], str(idx),
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
        )
        set_cell_text(row.cells[1], c["name"], bold=True)

        if c["linkedin"]:
            row.cells[2].text = ""
            p = row.cells[2].paragraphs[0]
            add_hyperlink(p, c["linkedin"], "Profile Link")
        else:
            set_cell_text(row.cells[2], "—")

        set_cell_text(row.cells[3], c["info"])

        flag = c.get("flag", "")
        flag_color = FLAG_COLORS.get(flag)
        flag_label = flag.capitalize() if flag else "—"
        set_cell_text(
            row.cells[4], flag_label,
            bold=True, color=flag_color,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
        )

    widths = [Cm(0.8), Cm(3.5), Cm(3), Cm(8), Cm(1.5)]
    for row in table.rows:
        for i, width in enumerate(widths):
            row.cells[i].width = width

    # Legend
    doc.add_paragraph("")
    doc.add_heading("Flag Legend", level=2)
    legend = doc.add_paragraph()
    run_g = legend.add_run("Green")
    run_g.bold = True
    run_g.font.color.rgb = FLAG_COLORS["green"]
    legend.add_run(
        " — Strong match, good candidate to approach\n"
    )
    run_y = legend.add_run("Yellow")
    run_y.bold = True
    run_y.font.color.rgb = FLAG_COLORS["yellow"]
    legend.add_run(
        " — Recently changed jobs or minor concerns "
        "(domestic company, unclear scope)\n"
    )
    run_r = legend.add_run("Red")
    run_r.bold = True
    run_r.font.color.rgb = FLAG_COLORS["red"]
    legend.add_run(
        " — Excluded (recruiter, founder, too senior age, "
        "below target seniority)"
    )

    doc.save(str(output))
    print(f"Saved: {output}")


if __name__ == "__main__":
    candidates = load_candidates(CSV_PATH)
    print(f"Found {len(candidates)} unique candidates after cleaning")
    for c in candidates:
        li = "LinkedIn" if c["linkedin"] else "No LinkedIn"
        print(f"  [{c['score']}] [{c['flag']}] {c['name']} ({li})")
    build_docx(candidates, DOCX_PATH)
